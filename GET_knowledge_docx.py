import requests
import json
from bs4 import BeautifulSoup, NavigableString, Tag
from datetime import datetime
import re
import os
import argparse
from docx import Document
from io import BytesIO
from PIL import Image
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from urllib.parse import urlencode
import html2text
import html
from html import unescape
import shutil
from html2docx import html2docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

load_dotenv()

def add_hyperlink(paragraph, url, text, styles=None):
    if styles is None:
        styles = {}

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    if styles.get("bold"):
        b = OxmlElement("w:b")
        rPr.append(b)
    if styles.get("italic"):
        i = OxmlElement("w:i")
        rPr.append(i)
    if styles.get("strike"):
        strike = OxmlElement("w:strike")
        rPr.append(strike)
    if styles.get("underline", True):
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if "color" in styles:
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), styles["color"])
        rPr.append(color_elem)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text.strip()
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def clean_inline_spans(html_content):
    if not html_content:
        return ""

    soup = BeautifulSoup(html_content, 'html.parser')

    def get_text_from_node(node):
        if isinstance(node, NavigableString):
            return str(node)
        elif isinstance(node, Tag):
            # Block-level tags that should create paragraph breaks
            block_tags = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr', 'table', 'section', 'article',
                          'br']

            if node.name == 'br':
                return '\n'  # treat line break

            if node.name in block_tags:
                # Join children with spaces, then add a line break after block
                inner_text = ''.join(get_text_from_node(c) for c in node.children)
                return inner_text.strip() + '\n\n'
            else:
                # Inline elements: join children with no added breaks, just spaces
                inner_text = ''.join(get_text_from_node(c) for c in node.children)
                return inner_text

        return ''

    text = get_text_from_node(soup)
    text = unescape(text)
    # Normalize whitespace: replace multiple spaces/newlines with single spaces except paragraph breaks
    # First replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Then replace multiple newlines with exactly two newlines (paragraph break)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    # Trim leading/trailing whitespace
    text = text.strip()

    return text


def get_bearer_token():
    url = "https://lendlease.service-now.com/oauth_token.do"
    payload_dict = {
        'grant_type': 'password',
        'username': os.getenv('SNOW_USERNAME'),
        'password': os.getenv('SNOW_PASSWORD'),
        'client_id': os.getenv('SNOW_CLIENT_ID'),
        'client_secret': os.getenv('SNOW_CLIENT_SECRET')
    }
    payload = urlencode(payload_dict)

    headers = {
        'Content-Type': 'application/x-www-form-urlencoded',
    }
    response = requests.post(url, data=payload, headers=headers)
    print(f"Response Status Code: {response.status_code}")
    data = response.json()
    return data['access_token']


def clean_html_text(html_content):
    """Convert HTML to clean text"""
    if not html_content:
        return ""

    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Convert to text while preserving some structure
    text = soup.get_text()

    # Clean up extra whitespace and line breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newline
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    text = text.strip()

    return text


def format_kb_article_backup(article):
    """Format a single knowledge base article for text backup"""
    formatted = []

    # Title/Number
    if article.get('number'):
        formatted.append(f"Article: {article['number']}")
        formatted.append("=" * 50)

    # Metadata
    if article.get('sys_created_on'):
        formatted.append(f"Created: {article['sys_created_on']}")
    if article.get('sys_updated_on'):
        formatted.append(f"Updated: {article['sys_updated_on']}")
    if article.get('workflow_state'):
        formatted.append(f"Status: {article['workflow_state']}")

    formatted.append("")  # Empty line

    # Main content
    if article.get('text'):
        clean_text = clean_html_text(article['text'])
        formatted.append("CONTENT:")
        formatted.append("-" * 20)
        formatted.append(clean_text)

    formatted.append("\n" + "=" * 80 + "\n")  # Separator between articles

    return "\n".join(formatted)


def format_kb_article_to_docx(doc, article):
    """Add a formatted knowledge base article to the Word document"""

    # Article title/number
    if article.get('number'):
        title = doc.add_heading(f"Article: {article['number']}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    meta_fields = [
        ("Created:", article.get('sys_created_on')),
        ("Updated:", article.get('sys_updated_on')),
        ("Updated By:", article.get('sys_updated_by')),
        ("Sys Domain:", article.get('sys_domain', {}).get('display_value') if isinstance(article.get('sys_domain'),
                                                                                         dict) else article.get(
            'sys_domain')),
        ("x_caukp_ebonding_no_return:", article.get('x_caukp_ebonding_no_return')),
        ("Select Portal:",
         article.get('u_select_portal', {}).get('display_value') if isinstance(article.get('u_select_portal'),
                                                                               dict) else article.get(
             'u_select_portal')),
        ("Status:", article.get('workflow_state')),
        ("Created By:", article.get('sys_created_by')),
        ("Published:", article.get('published')),
        ("Author:",
         article.get('author', {}).get('display_value') if isinstance(article.get('author'), dict) else article.get(
             'author')),
        ("x_caukp_ebonding_integration_mode:", article.get('x_caukp_ebonding_integration_mode')),
        ("Helpful Count:", article.get('helpful_count')),
        ("System Domain Path:", article.get('sys_domain_path')),
        ("View Count (All):", article.get('u_view_count_all')),
        ("Version:",
         article.get('version', {}).get('display_value') if isinstance(article.get('version'), dict) else article.get(
             'version')),
        ("Active:", article.get('active')),
        ("Topic:", article.get('topic')),
        ("Valid To:", article.get('valid_to')),
        ("KB Category:", article.get('kb_category', {}).get('display_value') if isinstance(article.get('kb_category'),
                                                                                           dict) else article.get(
            'kb_category')),
        ("Meta description:", article.get('meta_description')),
        ("KB Knowledge base:",
         article.get('kb_knowledge_base', {}).get('display_value') if isinstance(article.get('kb_knowledge_base'),
                                                                                 dict) else article.get(
             'kb_knowledge_base')),
        ("Meta:", article.get('meta')),
        ("U Problem Type:", article.get('u_problem_type')),
        ("Display number:", article.get('display_number')),
        ("Base version:",
         article.get('base_version', {}).get('display_value') if isinstance(article.get('base_version'),
                                                                            dict) else article.get('base_version')),
        ("Short description:", article.get('short_description')),
        ("Direct:", article.get('direct')),
        ("Disable suggesting:", article.get('disable_suggesting')),
        ("Class name:", article.get('sys_class_name')),
        ("Article Id:", article.get('article_id')),
        ("Sys Id:", article.get('sys_id')),
        ("Use Count:", article.get('use_count')),
        ("Flagged:", article.get('flagged')),
        ("Disable commenting:", article.get('disable_commenting')),
        ("Adding to home page:", article.get('u_add_to_homepage')),
        ("Display attachments:", article.get('display_attachments')),
        ("Latest:", article.get('latest')),
        ("Summary:",
         article.get('summary', {}).get('display_value') if isinstance(article.get('summary'), dict) else article.get(
             'summary')),
        ("Sys View Count:", article.get('sys_view_count')),
        ("Revised by:", article.get('revised_by', {}).get('display_value') if isinstance(article.get('revised_by'),
                                                                                         dict) else article.get(
            'revised_by')),
        ("Article Type:", article.get('article_type')),
        ("Needs review:", article.get('u_needs_review')),
        ("Sys Mod Count:", article.get('sys_mod_count')),
        ("View as allowed:", article.get('view_as_allowed')),
        ("Category:", article.get('category')),
        ("Reminder send date:", article.get('u_reminder_send_date')),
        ("Wiki:", article.get('wiki')),
        ("Rating:", article.get('rating')),
        ("Source:", article.get('source')),
        ("x_caukp_ebonding_sdc:", article.get('x_caukp_ebonding_sdc')),
        ("Scheduled Publish date:", article.get('scheduled_publish_date')),
        ("Image:", article.get('image')),
        ("KBI Uniqueid", article.get('u_kbi_uniqueid')),
        ("cmdb ci:", article.get('cmdb_ci')),
        ("Can Read User Criteria:", article.get('can_read_user_criteria')),
        ("Cannot Read User Criteria:", article.get('cannot_read_user_criteria')),
        ("x caukp ebonding requester id:", article.get('x_caukp_ebonding_requester_id')),
        ("Last Review date:", article.get('u_last_review_date')),
        ("x caukp ebonding provider id:", article.get('x_caukp_ebonding_provider_id')),
        ("Roles:", article.get('roles')),
        ("Description:", article.get('description')),
        ("sn_grc_target_table:", article.get('sn_grc_target_table')),
        ("Retired:", article.get('retired')),
        ("Video URL:", article.get('u_video_url')),
        ("sn_grc_source:", article.get('sn_grc_source')),
        ("Sys Tags:", article.get('sys_tags')),
        ("Replacement Article:", article.get('replacement_article')),
        ("x caukp ebonding provider:", article.get('x_caukp_ebonding_provider')),
        ("Taxonomy Topic:", article.get('taxonomy_topic')),
        ("x caukp ebonding requester:", article.get('x_caukp_ebonding_requester')),
        ("Ownership group:", article.get('ownership_group')),
    ]
    # Filter out None values and create table
    meta_fields = [item for item in meta_fields if item[1]]
    if meta_fields:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for label, value in meta_fields:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # Add space after table
        doc.add_paragraph()

    # Main content
    if article.get('text'):
        content_heading = doc.add_heading('Content', level=2)
        add_html_with_images(doc, article['text'])
        clean_text = clean_inline_spans(article['text'])  # clean_html_text(article['text'])

        # Split content into paragraphs and add them
        # paragraphs = clean_text.split('\n\n')
        # for para_text in paragraphs:
        #     if para_text.strip():
        #         doc.add_paragraph(para_text.strip())


def download_attachments_for_article(table_sys_id, output_dir, headers):
    """Download attachments for a specific KB article and save them in its folder,
    refresh token if 401 Unauthorized is received."""

    attachment_url = f"https://lendlease.service-now.com/api/now/attachment?sysparm_query=table_sys_id={table_sys_id}"

    def try_download(headers):
        try:
            response = requests.get(attachment_url, headers=headers)
            if response.status_code == 401:
                return 'unauthorized', None
            elif response.status_code != 200:
                print(f"❌ Failed to get attachment list for {table_sys_id}. Status code: {response.status_code}")
                return 'failed', None

            data = response.json()
            attachments = data.get('result', [])
            if not attachments:
                print(f"📎 No attachments found for {table_sys_id}")
                return 'empty', None

            print(f"📎 Found {len(attachments)} attachment(s) for {table_sys_id}")
            return 'success', attachments
        except Exception as e:
            print(f"❌ Exception while fetching attachments: {e}")
            return 'error', None

    status, attachments = try_download(headers)

    if status == 'unauthorized':
        print("🔄 Access token expired, refreshing token...")
        # Refresh token here and update headers
        new_token = get_bearer_token()
        headers['Authorization'] = f'Bearer {new_token}'
        # Retry once with new token
        status, attachments = try_download(headers)
        if status == 'unauthorized':
            print("❌ Token refresh failed or new token also unauthorized.")
            return
        elif status != 'success':
            return

    if status != 'success':
        return

    # Download each attachment
    for attachment in attachments:
        file_name = attachment.get('file_name')
        sys_id = attachment.get('sys_id')
        file_name = f"{sys_id}_{file_name}" if file_name else f"{table_sys_id}_attachment"
        download_link = attachment.get('download_link')
        file_size = attachment.get('size_bytes')

        if download_link and file_name:
            try:
                file_response = requests.get(download_link, headers=headers)
                if file_response.status_code == 200:
                    file_path = os.path.join(output_dir, file_name)
                    with open(file_path, 'wb') as f:
                        f.write(file_response.content)
                    print(f"   ✓ Downloaded: {file_name} ({file_size} bytes)")
                else:
                    print(f"   ✗ Failed to download {file_name} (Status {file_response.status_code})")
            except Exception as e:
                print(f"   ✗ Error downloading {file_name}: {e}")


def add_html_with_images(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    def is_inline(elem):
        # Tags usually inline in HTML
        inline_tags = {'span', 'a', 'b', 'i', 'u', 'em', 'strong', 'small', 'sub', 'sup', 'mark', 'code', 'br'}
        return elem.name in inline_tags if elem.name else False

    def process_element(elem, parent_paragraph=None):
        if isinstance(elem, NavigableString):
            # Append text to the paragraph if exists, else create new
            text = str(elem).strip()
            if text:
                if parent_paragraph is None:
                    parent_paragraph = doc.add_paragraph()
                parent_paragraph.add_run(text + ' ')
            return parent_paragraph

        elif elem.name == 'img':
            # Add image placeholder in a new paragraph
            src = elem.get('src', '')
            import re
            sysid_match = re.search(r'sys_id=([a-zA-Z0-9]+)', src)
            placeholder_text = "[IMAGE_PLACEHOLDER:UNKNOWN]"
            if sysid_match:
                sysid = sysid_match.group(1)
                placeholder_text = f"[IMAGE_PLACEHOLDER:{sysid}]"
            if parent_paragraph is not None:
                    # Insert placeholder run into existing paragraph (inside table cell)
                    parent_paragraph.add_run(placeholder_text + " ")
            else:
                # No paragraph context, create a new one
                para = doc.add_paragraph()
                para.add_run(placeholder_text)
            return parent_paragraph
        elif elem.name == 'table':
            add_html_table(doc, elem)
            return None
        elif elem.name == 'a':
            href = elem.get('href')
            link_text = elem.get_text(strip=True)
            if href and link_text:
                if parent_paragraph is None:
                    parent_paragraph = doc.add_paragraph()
                add_hyperlink(parent_paragraph, href, link_text)
            return parent_paragraph

        elif is_inline(elem):
            # NEW: Handle strikethrough spans
            if elem.name == 'span' and 'text-decoration: line-through' in elem.get('style', '').lower():
                text = elem.get_text().strip()
                if text:
                    if parent_paragraph is None:
                        parent_paragraph = doc.add_paragraph()
                    run = parent_paragraph.add_run(text + ' ')
                    run.font.strike = True
                return parent_paragraph
            
            # Existing inline handling
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
                
            for child in elem.children:
                parent_paragraph = process_element(child, parent_paragraph)
                
            return parent_paragraph

        else:
            # Block element handling
            for child in elem.children:
                process_element(child, None)
            return None

    # Process top-level elements
    top_level = soup.body.contents if soup.body else soup.contents
    for child in top_level:
        process_element(child, None)


def add_html_with_images(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    def is_inline(elem):
        inline_tags = {'span', 'a', 'b', 'i', 'u', 'em', 'strong', 'small', 'sub', 'sup', 'mark', 'code', 'br'}
        return elem.name in inline_tags if elem.name else False

    def process_element(elem, parent_paragraph=None):
        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                if parent_paragraph is None:
                    parent_paragraph = doc.add_paragraph()
                parent_paragraph.add_run(text + ' ')
            return parent_paragraph

        elif elem.name == 'img':
            # ... existing image handling code ...
            return None

        elif is_inline(elem):
            # NEW: Handle strikethrough spans
            if elem.name == 'span' and 'text-decoration: line-through' in elem.get('style', '').lower():
                text = elem.get_text().strip()
                if text:
                    if parent_paragraph is None:
                        parent_paragraph = doc.add_paragraph()
                    run = parent_paragraph.add_run(text + ' ')
                    run.font.strike = True
                return parent_paragraph
            
            # Existing inline handling
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
                
            for child in elem.children:
                parent_paragraph = process_element(child, parent_paragraph)
                
            return parent_paragraph

        else:
            # Block element handling
            for child in elem.children:
                process_element(child, None)
            return None

    # Process top-level elements
    top_level = soup.body.contents if soup.body else soup.contents
    for child in top_level:
        process_element(child, None)
        
def add_html_table(doc, table_elem):
        rows = table_elem.find_all('tr')
        if not rows:
            return
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'

        def process_cell_content(cell_elem, cell):
            # Clear any existing paragraphs in the cell (usually empty initially)
            for para in cell.paragraphs:
                p = para._element
                p.getparent().remove(p)
            # Add one new paragraph to the cell
            para = cell.add_paragraph()

            def process_element(elem, parent_paragraph):
                from bs4 import NavigableString, Tag
                if isinstance(elem, NavigableString):
                    text = str(elem).strip()
                    if text:
                        parent_paragraph.add_run(text + ' ')
                elif isinstance(elem, Tag):
                    if elem.name == 'img':
                        src = elem.get('src', '')
                        import re
                        sysid_match = re.search(r'sys_id=([a-zA-Z0-9]+)', src)
                        placeholder_text = "[IMAGE_PLACEHOLDER:UNKNOWN]"
                        if sysid_match:
                            sysid = sysid_match.group(1)
                            placeholder_text = f"[IMAGE_PLACEHOLDER:{sysid}]"
                        parent_paragraph.add_run(placeholder_text + ' ')
                    elif elem.name == 'br':
                        parent_paragraph.add_run('\n')
                    else:
                        for child in elem.children:
                            process_element(child, parent_paragraph)

            for child in cell_elem.children:
                process_element(child, para)

        for row_elem in rows:
            row_cells = row_elem.find_all(['td', 'th'])
            row = table.add_row()
            for i, cell_elem in enumerate(row_cells):
                process_cell_content(cell_elem, row.cells[i])


def replace_placeholders_with_images(docx_path, local_image_folder, output_path):
    doc = Document(docx_path)
    used_image_folder = os.path.join(local_image_folder, "used_images")
    os.makedirs(used_image_folder, exist_ok=True)
    used_sysids = set()

    def process_runs(runs, is_in_table=False):
        for run in runs:

            if "[IMAGE_PLACEHOLDER:" in run.text:
                match = re.search(r"\[IMAGE_PLACEHOLDER:([^\]]+)\]", run.text)
                if match:
                    sysid = match.group(1)
                    used_sysids.add(sysid)
                    found_file = None

                    # Search both main folder and used_images folder in case reused image was already moved
                    for search_dir in [local_image_folder, used_image_folder]:
                        for filename in os.listdir(search_dir):
                            if sysid in filename:
                                found_file = os.path.join(search_dir, filename)
                                break
                        if found_file:
                            break

                    if found_file:

                        run.clear()

                        with open(found_file, 'rb') as f:
                            img_stream = BytesIO(f.read())
                            width = Inches(0.5) if is_in_table else Inches(4)
                            # run.add_picture(img_stream, width=width)
                            try:
                                
                                img = Image.open(found_file)
                                buffer = BytesIO()
                                img.save(buffer, format=img.format)
                                buffer.seek(0)
                                run.add_picture(buffer, width=width)
                            except  Exception as e:
                                print(f"❌ Failed to add image for sys_id {sysid}. Reason: {e}")
                                run.text = f"[Failed to insert image: {os.path.basename(found_file)}]"
                                print("img_stream:", img_stream)
                                print(f"❌ Error adding image for sys_id {sysid}. Image stream might be empty or invalid.")
                                # If you want to force it, you could try to open the image with Pillow,
                                # but that would require a new function or more code than just a try-except.
                                # If you want to skip, just leave this block empty or put a comment.
                                pass

                    else:
                        run.text = f"[No image found for {sysid}]"

    # Process top-level paragraphs
    for paragraph in doc.paragraphs:
        process_runs(paragraph.runs, is_in_table=False)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs, is_in_table=True)

    doc.save(output_path)

    # Move used images (only once per sys_id)
    for sysid in used_sysids:
        for filename in os.listdir(local_image_folder):
            if sysid in filename:
                src = os.path.join(local_image_folder, filename)
                dst = os.path.join(used_image_folder, filename)
                if not os.path.exists(dst):  # Avoid overwriting if already moved
                    shutil.move(src, dst)
                break

# Parse command-line argument for knowledge base ID
parser = argparse.ArgumentParser(description='Download and export KB articles from ServiceNow')
parser.add_argument('kb_id', type=str, help='Knowledge Base sys_id (e.g., 01125e5a1b9b685017eeebd22a4bcb44)')
args = parser.parse_args()
kb_id = args.kb_id

# Your API call
#url = f"https://lendlease.service-now.com/api/now/table/kb_knowledge?sysparm_query=sys_class_name!=^publishedISNOTEMPTY^latest=true^kb_knowledge_base={kb_id}&sysparm_display_value=true"
url = f"https://lendlease.service-now.com/api/now/table/kb_knowledge?sysparm_query=workflow_state=published^active=true^latest=true^publishedISNOTEMPTY^kb_knowledge_base={kb_id}&sysparm_display_value=true"
payload = {}

token = get_bearer_token()
headers = {
    'Authorization': f'Bearer {token}',
    'Cookie': 'BIGipServerpool_lendlease=c5889ad29f701618e3baa37002034b82; JSESSIONID=3901AC59B602B51CE1CF74C8956FD362; glide_node_id_for_js=fc4812175032dd94c0ff92cf846b17cf27f0dce0a6beb49e12e5c7bb0f48d836; glide_session_store=6360D6592B3D6E50E412F41CD891BF5D; glide_user_activity=U0N2M18xOnRMdkppdFlTN2o2cFlnUVdaQ092UjZ6S0pFdXV0dmZBb3BMcGxVa0hrZ1E9OlVBQWc4QWozUERYQi9mVCs2WDRJa0hTRTgwQjkxMGZkMzUrNGxlUXRNUW89; glide_user_route=glide.5a07cc0a1b859ed021434a69d48daaeb'
}

try:
    # Make API request
    response = requests.request("GET", url, headers=headers, data=payload)

    if response.status_code == 200:
        # Parse JSON response
        data = response.json()

        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Create Word document
        doc = Document()

        # Document title and header
        title = doc.add_heading('Lendlease Knowledge Base Articles', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document info
        info_para = doc.add_paragraph()
        info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Total Articles: {len(data.get('result', []))}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add separator
        doc.add_paragraph("_" * 80)

        # Process each article
        articles = data.get('result', [])
        for i, article in enumerate(articles):
            # Generate a new document for each article
            doc = Document()

            # Document title
            title = doc.add_heading('Lendlease Knowledge Base Article', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Article info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            info_para.add_run(f"Article Number: {article.get('number', 'Unknown')}")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("_" * 80)

            # Add article content
            format_kb_article_to_docx(doc, article)

            # Save each article as a separate .docx file
            safe_article_number = re.sub(r'[^\w\-_. ]', '_', article.get('number', f"article_{i + 1}"))
            # Create output directory if it doesn't exist
            # output_dir = "KB_docx_files_{timestamp}".format(timestamp=timestamp)
            # os.makedirs(output_dir, exist_ok=True)
            # Create parent directory
            parent_dir = f"KB_docx_files_{timestamp}"
            os.makedirs(parent_dir, exist_ok=True)

            # Create subfolder for each article number
            article_number = article.get('number', f"article_{i + 1}")
            safe_article_number = re.sub(r'[^\w\-_. ]', '_', article_number)
            article_dir = os.path.join(parent_dir, safe_article_number)
            os.makedirs(article_dir, exist_ok=True)

            # Save the .docx file into the folder
            docx_filename = f"kb_article_{safe_article_number}.docx"
            docx_path = os.path.join(article_dir, docx_filename)
            doc.save(docx_path)
            print(f"📄 Saved: {docx_path}")

            # Download attachments for this article
            table_sys_id = article.get('sys_id')
            if table_sys_id:
                download_attachments_for_article(table_sys_id, article_dir, headers)
            else:
                print(f"⚠️ No sys_id found for article {safe_article_number}, skipping attachment download.")

            # After saving the article DOCX (e.g., article_docx_path)
            replace_placeholders_with_images(
                docx_path,
                local_image_folder=article_dir,  # Your attachment folder for the article
                output_path=os.path.join(article_dir, docx_filename)
            )
            print(f"📊 Processed Inline attachments for article {safe_article_number}")
        print(f"📊 Processed {len(articles)} articles")



    else:
        print(f"❌ API request failed with status code: {response.status_code}")
        print(f"Response: {response.text[:200]}...")

except requests.exceptions.RequestException as e:
    print(f"❌ Request failed: {e}")
except json.JSONDecodeError as e:
    print(f"❌ Failed to parse JSON response: {e}")
except Exception as e:
    print(f"❌ An error occurred: {e}")
