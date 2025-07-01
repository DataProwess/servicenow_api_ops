import requests
import json
from bs4 import BeautifulSoup, NavigableString, Tag
from datetime import datetime
import re
import os
import argparse
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from urllib.parse import urlencode
import html2text
import html
from html import unescape
from html2docx import html2docx
import base64
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

load_dotenv()

 

def add_hyperlink(paragraph, url, text, styles=None):
    if styles is None:
        styles = {}

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    if styles.get("bold"):
        b = OxmlElement("w:b")
        rPr.append(b)
    if styles.get("italic"):
        i = OxmlElement("w:i")
        rPr.append(i)
    if styles.get("strike"):
        strike = OxmlElement("w:strike")
        rPr.append(strike)
    if styles.get("underline", True):
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if "color" in styles:
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), styles["color"])
        rPr.append(color_elem)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text.strip()
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph
 

def replace_img_with_confluence_macro(html_content, attachments):
    """Replace <img> tags with Confluence <ac:image> macros referencing attachment filenames."""
    soup = BeautifulSoup(html_content, "html.parser")
    # Build a set of attachment filenames for quick lookup (strip sys_id prefix if present)
    attachment_filenames = set()
    for a in attachments:
        fname = a['file_name']
        # Remove sys_id prefix if present (e.g., 'abcdef12345_image.png' -> 'image.png')
        if '_' in fname and len(fname.split('_')[0]) == 32:
            fname_noid = '_'.join(fname.split('_')[1:])
            attachment_filenames.add(fname_noid)
            attachment_filenames.add(fname)
        else:
            attachment_filenames.add(fname)

    for img in soup.find_all("img"):
        src = img.get("src", "")
        filename = None
        # Try to extract filename from src
        if 'sys_id=' in src:
            # ServiceNow attachment URL: ...?sys_id=abcdef12345
            sysid = src.split('sys_id=')[-1].split('&')[0]
            # Try to match sysid to attachments
            for a in attachments:
                if sysid in a['file_name'] or sysid == a.get('sys_id'):
                    filename = a['file_name']
                    break
        if not filename:
            # Fallback: use last segment of src
            filename = src.split('/')[-1].split('?')[0]
            # Try to match filename
            for f in attachment_filenames:
                if filename in f:
                    filename = f
                    break

        if filename in attachment_filenames:
            # Build the Confluence image macro
            macro = BeautifulSoup(
                f'<ac:image><ri:attachment ri:filename="{filename}"/></ac:image>',
                "html.parser"
            )
            img.replace_with(macro)
        else:
            # Remove img tags that don't match an uploaded attachment
            img.decompose()
    return str(soup)



def clean_inline_spans(html_content):
    if not html_content:
        return ""

    soup = BeautifulSoup(html_content, 'html.parser')

    def get_text_from_node(node):
        if isinstance(node, NavigableString):
            return str(node)
        elif isinstance(node, Tag):
            # Block-level tags that should create paragraph breaks
            block_tags = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr', 'table', 'section', 'article', 'br']

            if node.name == 'br':
                return '\n'  # treat line break

            if node.name in block_tags:
                # Join children with spaces, then add a line break after block
                inner_text = ''.join(get_text_from_node(c) for c in node.children)
                return inner_text.strip() + '\n\n'
            else:
                # Inline elements: join children with no added breaks, just spaces
                inner_text = ''.join(get_text_from_node(c) for c in node.children)
                return inner_text

        return ''

    text = get_text_from_node(soup)
    text = unescape(text)
    # Normalize whitespace: replace multiple spaces/newlines with single spaces except paragraph breaks
    # First replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Then replace multiple newlines with exactly two newlines (paragraph break)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    # Trim leading/trailing whitespace
    text = text.strip()

    return text

def get_bearer_token():
    url = "https://lendlease.service-now.com/oauth_token.do"
    payload_dict = {
        'grant_type': 'password',
        'username': os.getenv('SNOW_USERNAME'),
        'password': os.getenv('SNOW_PASSWORD'),
        'client_id': os.getenv('SNOW_CLIENT_ID'),
        'client_secret': os.getenv('SNOW_CLIENT_SECRET')
    }
    payload = urlencode(payload_dict)

    headers = {
    'Content-Type': 'application/x-www-form-urlencoded',
    }
    response = requests.post(url, data=payload, headers=headers)
    print(f"Response Status Code: {response.status_code}")
    data = response.json()
    return data['access_token']

def clean_html_text(html_content):
    """Convert HTML to clean, linear text without unnecessary line breaks from inline tags like <span>."""
    if not html_content:
        return ""
    
    soup = BeautifulSoup(html_content, 'html.parser')

    # Join all text chunks ignoring artificial breaks between inline tags
    text = ''.join(soup.stripped_strings)

    # Optional: unescape HTML entities like &lsquo; and &rsquo;
    from html import unescape
    text = unescape(text)

    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    return text


def download_attachments_for_article(table_sys_id, output_dir, headers):
    """Download attachments for a specific KB article and save them in its folder,
    refresh token if 401 Unauthorized is received."""

    attachment_url = f"https://lendlease.service-now.com/api/now/attachment?sysparm_query=table_sys_id={table_sys_id}"

    def try_download(headers):
        try:
            response = requests.get(attachment_url, headers=headers)
            if response.status_code == 401:
                return 'unauthorized', None
            elif response.status_code != 200:
                print(f"❌ Failed to get attachment list for {table_sys_id}. Status code: {response.status_code}")
                return 'failed', None
            
            data = response.json()
            attachments = data.get('result', [])
            if not attachments:
                print(f"📎 No attachments found for {table_sys_id}")
                return 'empty', None
            
            print(f"📎 Found {len(attachments)} attachment(s) for {table_sys_id}")
            return 'success', attachments
        except Exception as e:
            print(f"❌ Exception while fetching attachments: {e}")
            return 'error', None

    status, attachments = try_download(headers)

    if status == 'unauthorized':
        print("🔄 Access token expired, refreshing token...")
        # Refresh token here and update headers
        new_token = get_bearer_token()
        headers['Authorization'] = f'Bearer {new_token}'
        # Retry once with new token
        status, attachments = try_download(headers)
        if status == 'unauthorized':
            print("❌ Token refresh failed or new token also unauthorized.")
            return []
        elif status != 'success':
            return []

    if status != 'success':
        return []

    downloaded_attachments = []
    # Download each attachment
    for attachment in attachments:
        file_name = attachment.get('file_name')
        sys_id = attachment.get('sys_id')
        file_name = f"{sys_id}_{file_name}" if file_name else f"{table_sys_id}_attachment"
        download_link = attachment.get('download_link')
        file_size = attachment.get('size_bytes')

        if download_link and file_name:
            try:
                file_response = requests.get(download_link, headers=headers)
                if file_response.status_code == 200:
                    file_path = os.path.join(output_dir, file_name)
                    with open(file_path, 'wb') as f:
                        f.write(file_response.content)
                    print(f"   ✓ Downloaded: {file_name} ({file_size} bytes)")
                    downloaded_attachments.append({
                        'file_name': file_name,
                        'file_path': file_path,
                        'sys_id': sys_id,
                        'size_bytes': file_size
                    })
                else:
                    print(f"   ✗ Failed to download {file_name} (Status {file_response.status_code})")
            except Exception as e:
                print(f"   ✗ Error downloading {file_name}: {e}")
    
    return downloaded_attachments

        
def add_html_table(doc, table_elem):
        rows = table_elem.find_all('tr')
        if not rows:
            return
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'

        def process_cell_content(cell_elem, cell):
            # Clear any existing paragraphs in the cell (usually empty initially)
            for para in cell.paragraphs:
                p = para._element
                p.getparent().remove(p)
            # Add one new paragraph to the cell
            para = cell.add_paragraph()

            def process_element(elem, parent_paragraph):
                from bs4 import NavigableString, Tag
                if isinstance(elem, NavigableString):
                    text = str(elem).strip()
                    if text:
                        parent_paragraph.add_run(text + ' ')
                elif isinstance(elem, Tag):
                    if elem.name == 'img':
                        src = elem.get('src', '')
                        import re
                        sysid_match = re.search(r'sys_id=([a-zA-Z0-9]+)', src)
                        placeholder_text = "[IMAGE_PLACEHOLDER:UNKNOWN]"
                        if sysid_match:
                            sysid = sysid_match.group(1)
                            placeholder_text = f"[IMAGE_PLACEHOLDER:{sysid}]"
                        parent_paragraph.add_run(placeholder_text + ' ')
                    elif elem.name == 'br':
                        parent_paragraph.add_run('\n')
                    else:
                        for child in elem.children:
                            process_element(child, parent_paragraph)

            for child in cell_elem.children:
                process_element(child, para)

        for row_elem in rows:
            row_cells = row_elem.find_all(['td', 'th'])
            row = table.add_row()
            for i, cell_elem in enumerate(row_cells):
                process_cell_content(cell_elem, row.cells[i])


def add_html_with_images(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    def is_inline(elem):
        # Tags usually inline in HTML
        inline_tags = {'span', 'a', 'b', 'i', 'u', 'em', 'strong', 'small', 'sub', 'sup', 'mark', 'code', 'br'}
        return elem.name in inline_tags if elem.name else False

    def process_element(elem, parent_paragraph=None):
        if isinstance(elem, NavigableString):
            # Append text to the paragraph if exists, else create new
            text = str(elem).strip()
            if text:
                if parent_paragraph is None:
                    parent_paragraph = doc.add_paragraph()
                parent_paragraph.add_run(text + ' ')
            return parent_paragraph

        elif elem.name == 'img':
            # Add image placeholder in a new paragraph
            src = elem.get('src', '')
            import re
            sysid_match = re.search(r'sys_id=([a-zA-Z0-9]+)', src)
            placeholder_text = "[IMAGE_PLACEHOLDER:UNKNOWN]"
            if sysid_match:
                sysid = sysid_match.group(1)
                placeholder_text = f"[IMAGE_PLACEHOLDER:{sysid}]"
            if parent_paragraph is not None:
                    # Insert placeholder run into existing paragraph (inside table cell)
                    parent_paragraph.add_run(placeholder_text + " ")
            else:
                # No paragraph context, create a new one
                para = doc.add_paragraph()
                para.add_run(placeholder_text)
            return parent_paragraph
        elif elem.name == 'table':
            add_html_table(doc, elem)
            return None
        elif elem.name == 'a':
            href = elem.get('href')
            link_text = elem.get_text(strip=True)
            if href and link_text:
                if parent_paragraph is None:
                    parent_paragraph = doc.add_paragraph()
                add_hyperlink(parent_paragraph, href, link_text)
            return parent_paragraph

        elif is_inline(elem):
            # NEW: Handle strikethrough spans
            if elem.name == 'span' and 'text-decoration: line-through' in elem.get('style', '').lower():
                text = elem.get_text().strip()
                if text:
                    if parent_paragraph is None:
                        parent_paragraph = doc.add_paragraph()
                    run = parent_paragraph.add_run(text + ' ')
                    run.font.strike = True
                return parent_paragraph
            
            # Existing inline handling
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
                
            for child in elem.children:
                parent_paragraph = process_element(child, parent_paragraph)
                
            return parent_paragraph

        else:
            # Block element handling
            for child in elem.children:
                process_element(child, None)
            return None

    # Process top-level elements
    top_level = soup.body.contents if soup.body else soup.contents
    for child in top_level:
        process_element(child, None)



# NEW CONFLUENCE FUNCTIONS
def upload_attachment_to_confluence(confluence_url, username, api_token, page_id, file_path, file_name):
    """Upload an attachment to a Confluence page"""
    url = f"{confluence_url}/rest/api/content/{page_id}/child/attachment"
    
    headers = {
        'X-Atlassian-Token': 'no-check'
    }
    
    with open(file_path, 'rb') as f:
        files = {
            'file': (file_name, f, 'application/octet-stream')
        }
        
        response = requests.post(
            url,
            headers=headers,
            files=files,
            auth=(username, api_token)
        )
    
    if response.status_code == 200:
        attachment_data = response.json()
        return attachment_data['results'][0] if attachment_data.get('results') else None
    else:
        print(f"❌ Failed to upload attachment {file_name}: {response.status_code} - {response.text}")
        return None



def create_confluence_content(article, attachments):
    """Generate Confluence storage format content from KB article JSON"""
    
    # Start with title
    # content = f"<h1>{article.get('number', 'KB Article')}</h1>"
    
    # Add metadata as a table
    content = "<h2>Article Information</h2>"
    content += "<table><tbody>"
    
    # Key metadata fields
    important_fields = [
        ("Article Number", article.get('number')),
        ("KB Category:", article.get('kb_category', {}).get('display_value') if isinstance(article.get('kb_category'), dict) else article.get('kb_category')),
        # ("Meta description:",article.get('meta_description')),
        ("KB Knowledge base:", article.get('kb_knowledge_base', {}).get('display_value') if isinstance(article.get('kb_knowledge_base'), dict) else article.get('kb_knowledge_base')),
        ("Meta:",article.get('meta')),
    ]
    
    for label, value in important_fields:
        if value:
            content += f"<tr><td><strong>{label}</strong></td><td>{value}</td></tr>"
    
    content += "</tbody></table>"
    
    # Add main content
    if article.get('text'):
        content += "<h2>Content</h2>"
        # Clean up the HTML content for Confluence
        article_content = article.get('text', '')
        article_content = replace_img_with_confluence_macro(article_content, attachments)
        # Insert code to add video macro for each video attachment
        for attachment in attachments:
            if attachment['file_name'].lower().endswith(('.mp4', '.mov', '.avi', '.wmv')):
                video_macro = f"""
        <ac:structured-macro ac:name="multimedia">
        <ac:parameter ac:name="name">
            <ri:attachment ri:filename="{attachment['file_name']}"/>
        </ac:parameter>
        <ac:parameter ac:name="width">500</ac:parameter>
        <ac:parameter ac:name="autostart">false</ac:parameter>
        </ac:structured-macro>"""
        # Insert at the end of article_content, or at a specific location if you know where
                article_content += video_macro

        # Convert ServiceNow specific tags or clean up if needed
        content += article_content
    
    # Add attachments section if any
    if attachments:
        content += "<h2>Attachments</h2>"
        content += "<ul>"
        for attachment in attachments:
            content += f"<li><ac:link><ri:attachment ri:filename=\"{attachment['file_name']}\"/></ac:link></li>"
        content += "</ul>"
    
    return content

def create_or_update_confluence_page(confluence_url, username, api_token, space_key, article, attachments):
    """Create or update a Confluence page with KB article content"""
    
    page_title = f"{article.get('short_description')}"
    
    # Check if page already exists
    search_url = f"{confluence_url}/rest/api/content"
    search_params = {
        'title': page_title,
        'spaceKey': space_key,
        'expand': 'version'
    }
    
    response = requests.get(search_url, params=search_params, auth=(username, api_token))
    
    if response.status_code != 200:
        print(f"❌ Failed to search for existing page: {response.status_code}")
        return None
    
    search_results = response.json()
    existing_page = search_results['results'][0] if search_results['results'] else None
    
    # Generate content
    content = create_confluence_content(article, attachments)
    
    if existing_page:
        # Update existing page
        page_id = existing_page['id']
        current_version = existing_page['version']['number']
        
        update_data = {
            "version": {
                "number": current_version + 1
            },
            "title": page_title,
            "type": "page",
            "body": {
                "storage": {
                    "value": content,
                    "representation": "storage"
                }
            }
        }
        
        update_url = f"{confluence_url}/rest/api/content/{page_id}"
        response = requests.put(
            update_url,
            json=update_data,
            headers={'Content-Type': 'application/json'},
            auth=(username, api_token)
        )
        
        if response.status_code == 200:
            print(f"✅ Updated Confluence page: {page_title}")
            return response.json()
        else:
            print(f"❌ Failed to update page: {response.status_code} - {response.text}")
            return None
    
    else:
        # Create new page
        create_data = {
            "type": "page",
            "title": page_title,
            "space": {
                "key": space_key
            },
            "body": {
                "storage": {
                    "value": content,
                    "representation": "storage"
                }
            }
        }
        
        create_url = f"{confluence_url}/rest/api/content"
        response = requests.post(
            create_url,
            json=create_data,
            headers={'Content-Type': 'application/json'},
            auth=(username, api_token)
        )
        
        if response.status_code == 200:
            print(f"✅ Created new Confluence page: {page_title}")
            return response.json()
        else:
            print(f"❌ Failed to create page: {response.status_code} - {response.text}")
            return None

# Parse command-line arguments
parser = argparse.ArgumentParser(description='Download and export a specific KB article from ServiceNow to DOCX and Confluence')
parser.add_argument('article_number', type=str, help='KB article number (e.g., KB0020129)')
args = parser.parse_args()

article_number = args.article_number

# Get Confluence parameters from environment variables
confluence_url = os.getenv('CONFLUENCE_URL')
confluence_username = os.getenv('CONFLUENCE_USERNAME')
confluence_token = os.getenv('CONFLUENCE_TOKEN')
confluence_space = os.getenv('CONFLUENCE_SPACE')

# Updated API call to get only one article by number
url = f"https://lendlease.service-now.com/api/now/table/kb_knowledge?sysparm_query=number={article_number}^latest=true&sysparm_display_value=true"

# url = f"https://lendlease.service-now.com/api/now/table/kb_knowledge?sysparm_query=number={article_number}&sysparm_display_value=true"
# url = f"https://lendlease.service-now.com/api/now/table/kb_knowledge?sysparm_query=sys_class_name!=^publishedISNOTEMPTY^latest=true^number={article_number}&sysparm_display_value=true"
# url = f"https://lendlease.service-now.com/kb_view.do?sysparm_article={article_number}"
payload = {}

token = get_bearer_token()
headers = {
  'Authorization': f'Bearer {token}',
  'Cookie': 'BIGipServerpool_lendlease=c5889ad29f701618e3baa37002034b82; JSESSIONID=3901AC59B602B51CE1CF74C8956FD362; glide_node_id_for_js=fc4812175032dd94c0ff92cf846b17cf27f0dce0a6beb49e12e5c7bb0f48d836; glide_session_store=6360D6592B3D6E50E412F41CD891BF5D; glide_user_activity=U0N2M18xOnRMdkppdFlTN2o2cFlnUVdaQ082UjZ6S0pFdXV0dmZBb3BMcGxVa0hrZ1E9OlVBQWc4QWozUERYQi9mVCs2WDRJa0hTRTgwQjkxMGZkMzUrNGxlUXRNUW89; glide_user_route=glide.5a07cc0a1b859ed021434a69d48daaeb'
}
response = requests.get(url, headers=headers)

if response.status_code != 200:
    print(f"❌ Failed to fetch article {article_number}. Status code: {response.status_code}")
    exit(1)

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
data = response.json()
articles = data.get('result', [])

if not articles:
    print(f"❌ No article found with number {article_number}")
    exit(1)

article = articles[0]  # Just one
parent_dir = f"KB_docx_files_{timestamp}"
output_dir = os.path.join(parent_dir, article_number)
os.makedirs(output_dir, exist_ok=True)

# Download attachments
downloaded_attachments = download_attachments_for_article(article['sys_id'], output_dir, headers)


# Upload to Confluence if parameters are available in environment
if confluence_url and confluence_username and confluence_token and confluence_space:
    print("🚀 Uploading to Confluence...")
    
    # Create or update the Confluence page
    confluence_page = create_or_update_confluence_page(
        confluence_url,
        confluence_username,
        confluence_token,
        confluence_space,
        article,
        downloaded_attachments
    )
    
    if confluence_page:
        page_id = confluence_page['id']
        
        # Upload attachments to Confluence
        for attachment in downloaded_attachments:
            print(f"📎 Uploading attachment: {attachment['file_name']}")
            uploaded = upload_attachment_to_confluence(
                confluence_url,
                confluence_username,
                confluence_token,
                page_id,
                attachment['file_path'],
                attachment['file_name']
            )
            if uploaded:
                print(f"   ✅ Attachment uploaded successfully")
            else:
                print(f"   ❌ Failed to upload attachment")
        
        page_url = f"{confluence_url}/pages/viewpage.action?pageId={page_id}"
        print(f"✅ Confluence page available at: {page_url}")
    else:
        print("❌ Failed to create/update Confluence page")
        
else:
    missing_vars = []
    if not confluence_url: missing_vars.append('CONFLUENCE_URL')
    if not confluence_username: missing_vars.append('CONFLUENCE_USERNAME')
    if not confluence_token: missing_vars.append('CONFLUENCE_TOKEN')
    if not confluence_space: missing_vars.append('CONFLUENCE_SPACE')
    
    if missing_vars:
        print(f"ℹ️  Missing Confluence environment variables: {', '.join(missing_vars)}")
        print("   Add these to your .env file to enable Confluence upload.")
    else:
        print("ℹ️  No Confluence parameters provided. Skipping Confluence upload.")